"""
Script tạo báo cáo vận hành M-BFP Minimum Flow Valve Oscillation
Van Phong 1 BOT Thermal Power Plant
"""

from docx import Document
from docx.shared import Pt, Inches, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

def set_cell_background(cell, color_hex):
    """Set background color for table cell"""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def set_col_width(table, col_idx, width_cm):
    for row in table.rows:
        row.cells[col_idx].width = Cm(width_cm)

def add_heading(doc, text, level=1, color=None):
    p = doc.add_heading(text, level=level)
    if color:
        for run in p.runs:
            run.font.color.rgb = RGBColor(*color)
    return p

def add_para(doc, text, bold=False, indent=False):
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(11)
    return p

def add_table_header_row(table, headers, bg='1F4E79'):
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_background(cell, bg)

def add_table_row(table, values, bg=None, bold_first=False):
    row = table.add_row()
    for i, v in enumerate(values):
        cell = row.cells[i]
        cell.text = ''
        run = cell.paragraphs[0].add_run(str(v))
        run.font.size = Pt(10)
        if i == 0 and bold_first:
            run.bold = True
        if bg:
            set_cell_background(cell, bg)
    return row

# ─── Tạo document ────────────────────────────────────────────────────────────
doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin = Cm(3.0)
    section.right_margin = Cm(2.5)

# ─── HEADER ──────────────────────────────────────────────────────────────────
title_para = doc.add_paragraph()
title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title_para.add_run('VAN PHONG 1 BOT THERMAL POWER PLANT')
run.bold = True
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

title_para2 = doc.add_paragraph()
title_para2.alignment = WD_ALIGN_PARAGRAPH.CENTER
run2 = title_para2.add_run('OPERATION REPORT')
run2.bold = True
run2.font.size = Pt(13)
run2.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

doc.add_paragraph()

# Doc info table
info_table = doc.add_table(rows=5, cols=4)
info_table.style = 'Table Grid'
info_data = [
    ('Document No.', 'KV-OP-26-0001', 'Date', '13 June 2026'),
    ('Equipment', 'Motor-Driven Boiler Feed Pump (M-BFP)', 'Unit', 'Unit 1'),
    ('KKS Tag', '10LAC30AP010 KP01', 'Discipline', 'Operation'),
    ('Subject', 'Oscillation of M-BFP Minimum Flow Control Valve\n(10LAB33AA502)', 'Prepared by', 'Operations Dept.'),
    ('Rev.', 'A', 'Approved by', ''),
]
for r_idx, row_data in enumerate(info_data):
    row = info_table.rows[r_idx]
    for c_idx, val in enumerate(row_data):
        cell = row.cells[c_idx]
        cell.text = ''
        run = cell.paragraphs[0].add_run(val)
        run.font.size = Pt(10)
        if c_idx % 2 == 0:
            run.bold = True
            set_cell_background(cell, 'BDD7EE')
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if c_idx % 2 == 1 else WD_ALIGN_PARAGRAPH.LEFT

doc.add_paragraph()

# ─── 1. INTRODUCTION ─────────────────────────────────────────────────────────
add_heading(doc, '1. INTRODUCTION', level=2, color=(0x1F, 0x4E, 0x79))

intro_text = (
    "This report describes an abnormal operating event involving the Motor-Driven Boiler Feed Pump "
    "(M-BFP, KKS: 10LAC30AP010 KP01) at Van Phong 1 BOT Thermal Power Plant, Unit 1. "
    "During plant operation at partial load, the M-BFP Minimum Flow Control Valve "
    '(10LAB33AA502, hereafter "MIN FCV") was observed to oscillate repeatedly between '
    "open and closed states, causing unstable pump discharge pressure and increased mechanical "
    "stress on the valve actuator and associated piping.\n\n"
    "The M-BFP is a KSB CHTD 5/5 high-pressure barrel-type centrifugal pump rated at 915 m³/h "
    "at design point (85% BMCR + 5% margin). The MIN FCV (10LAB33AA502) is controlled by "
    "DCS logic sheet MD109 and serves to protect the pump against operating below minimum "
    "continuous stable flow (454 m³/h at design, 320 m³/h at 35% load) by recirculating "
    "water back to the Deaerator and Feed Water Tank."
)
p = doc.add_paragraph()
p.add_run(intro_text).font.size = Pt(11)

doc.add_paragraph()

# ─── 2. ABNORMAL STATUS ──────────────────────────────────────────────────────
add_heading(doc, '2. ABNORMAL STATUS', level=2, color=(0x1F, 0x4E, 0x79))

p = doc.add_paragraph()
p.add_run(
    "The following abnormal conditions were observed during the event:"
).font.size = Pt(11)

# Timeline table
doc.add_paragraph()
tl_table = doc.add_table(rows=1, cols=3)
tl_table.style = 'Table Grid'
tl_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(tl_table, ['Time', 'Parameter / Event', 'Observed Value / Status'])

timeline = [
    ('T+00:00', 'Unit load at partial load (~35–40% MCR)', '~250–270 MW'),
    ('T+00:00', 'M-BFP Inlet Water Flow (10LAB17CF301)', '~320–350 t/h'),
    ('T+00:00', 'M-BFP MIN FCV (10LAB33AA502) position', 'Cycling: 0% ↔ 100%'),
    ('T+00:05', 'M-BFP Outlet Pressure (10LAB33EE101 area)', 'Fluctuating ±5–8 bar'),
    ('T+00:10', 'M-BFP Bearing Housing Vibration', 'Elevated, ~3.2 mm/s'),
    ('T+00:15', 'DCS Alarm — MBFP MIN FCV ABN', 'ALARM active'),
    ('T+00:30', 'Operator intervened — switched MIN FCV to MANUAL', 'MAN mode, valve ~30% OPEN'),
    ('T+00:35', 'M-BFP flow and pressure stabilized', 'Flow ~360 t/h stable'),
    ('T+01:00', 'Plant engineer notified, investigation commenced', '—'),
]
for t, e, v in timeline:
    add_table_row(tl_table, [t, e, v])

doc.add_paragraph()
p = doc.add_paragraph()
p.add_run(
    "The oscillation frequency was approximately 1 cycle every 20–30 seconds. Each cycle caused "
    "a visible pressure spike on the M-BFP outlet pressure trend and hydrodynamic noise in the "
    "recirculation line. The event did not result in pump trip or loss of feedwater supply."
).font.size = Pt(11)

doc.add_paragraph()

# ─── 3. ANALYSIS ─────────────────────────────────────────────────────────────
add_heading(doc, '3. ANALYSIS', level=2, color=(0x1F, 0x4E, 0x79))

add_heading(doc, '3.1 Operating Context', level=3)
p = doc.add_paragraph()
p.add_run(
    "At the time of the incident, Unit 1 was operating at partial load. Based on data from "
    "the MDBFP Data Sheet (Doc: VP1-C-L2-M-LAC-00015), the M-BFP minimum continuous flow "
    "requirement is:"
).font.size = Pt(11)

spec_table = doc.add_table(rows=1, cols=3)
spec_table.style = 'Table Grid'
spec_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(spec_table, ['Operating Case', 'Minimum Flow (m³/h)', 'Pump Speed (rpm)'])
spec_data = [
    ('Design Point (85% BMCR + 5% margin)', '454', '6,160'),
    ('35% Load + 10% Margin', '320', '4,393'),
    ('Minimum recommended continuous speed', '300 (min.)', '1,800 (min.)'),
]
for row in spec_data:
    add_table_row(spec_table, row)

doc.add_paragraph()

add_heading(doc, '3.2 Control Logic Analysis (MD109 — M-BFP MIN FCV)', level=3)
p = doc.add_paragraph()
p.add_run(
    "The M-BFP MIN FCV (10LAB33AA502) is controlled by DCS logic block MD109 "
    "(Logic Diagrams for Boiler Feedwater Pump, Doc: VP1-C-L2-I-CB-00006, Sheet TM184/TM184B). "
    "The control logic operates as follows:"
).font.size = Pt(11)

bullets = [
    "The valve opens (recirculates to Deaerator) when MBFP Inlet Water Flow (10LAB17CF301) drops below the LOW setpoint.",
    "The valve closes (stops recirculation) when flow rises above the HIGH setpoint.",
    "The control employs a hysteresis band to prevent rapid cycling. From the logic diagram, the switching thresholds are configured with separate RESET (R=) and SET (S=) values.",
    "During AUTO mode, the valve position is modulated. During the event, the flow signal was near the control threshold (~320–360 t/h), causing the logic to repetitively switch the valve between OPEN and CLOSE commands.",
]
for b in bullets:
    p2 = doc.add_paragraph(style='List Bullet')
    p2.add_run(b).font.size = Pt(11)

doc.add_paragraph()

add_heading(doc, '3.3 Root Cause Analysis', level=3)
p = doc.add_paragraph()
p.add_run(
    "The root cause is identified as a combination of the following factors:"
).font.size = Pt(11)

causes = [
    ('RC-1', 'Insufficient PID hysteresis band',
     'The deadband (hysteresis) between the OPEN and CLOSE setpoints of the MIN FCV '
     'control logic (MD109) was narrower than the natural oscillation amplitude of the '
     'MBFP inlet flow signal. When flow fluctuated around the setpoint threshold, the '
     'logic alternately commanded the valve to open and close in rapid succession.'),
    ('RC-2', 'Flow signal noise / measurement instability',
     'The MBFP Inlet Water Flow transmitter (10LAB17CF301) may have exhibited signal '
     'fluctuation or noise at low-flow conditions (~320–360 t/h), further aggravating '
     'the control loop instability at the switching boundary.'),
    ('RC-3', 'No ramp rate limiting on MIN FCV at AUTO↔MANUAL transition',
     'Per the commissioning procedure test (Doc: VP1-C-L2-F-GEN-00029), the control '
     'specification states that if deviation between valve position and MV value exceeds '
     '0.6%, the MV ramp rate must be applied. Absence of adequate ramp limiting allowed '
     'full stroke operation on each cycle, increasing mechanical impact on the valve.'),
    ('RC-4', 'Plant operating near the minimum flow boundary',
     'At ~35% plant load, the M-BFP was operating near the minimum recommended flow '
     '(320 m³/h per data sheet). Small changes in load demand caused the actual flow '
     'to oscillate above and below the MIN FCV switching setpoint.'),
]

rc_table = doc.add_table(rows=1, cols=3)
rc_table.style = 'Table Grid'
rc_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(rc_table, ['Code', 'Cause', 'Description'])
for code, cause, desc in causes:
    row = rc_table.add_row()
    row.cells[0].text = code
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[0].paragraphs[0].runs[0].bold = True
    set_cell_background(row.cells[0], 'FFF2CC')
    row.cells[1].text = cause
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].paragraphs[0].runs[0].bold = True
    row.cells[2].text = desc
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)

doc.add_paragraph()

# ─── 4. SUGGESTION ───────────────────────────────────────────────────────────
add_heading(doc, '4. SUGGESTION', level=2, color=(0x1F, 0x4E, 0x79))

suggestions = [
    ('S-1', 'Review and widen MIN FCV control hysteresis band',
     'Evaluate the current OPEN/CLOSE setpoints in MD109 logic. Increase the hysteresis '
     'band (difference between OPEN setpoint and CLOSE setpoint) to at least 30–50 t/h '
     'to prevent oscillation when flow is near the threshold. Coordinate with Toshiba/KSB '
     'for recommended setpoint values.'),
    ('S-2', 'Apply signal filtering/damping on flow transmitter 10LAB17CF301',
     'Introduce a first-order low-pass filter or increase signal damping time constant '
     'on the MBFP Inlet Water Flow measurement to reduce high-frequency noise that '
     'triggers unnecessary valve cycling.'),
    ('S-3', 'Add MIN FCV ramp rate limiter',
     'Configure an adequate ramp rate (e.g., 5–10%/s) for the MIN FCV movement in AUTO '
     'mode to reduce mechanical shock on the valve disc, seat, and actuator during each '
     'open/close cycle. This is aligned with the commissioning requirement in '
     'Doc: VP1-C-L2-F-GEN-00029.'),
    ('S-4', 'Inspect MIN FCV actuator and valve internals',
     'Per KSB O&M Manual (Doc: 02.05.08, Section 7), the minimum flow valve should be '
     'checked during pump set inspection. After the oscillation event, visually inspect '
     'the actuator, positioner feedback, and valve disc/seat for wear or damage caused '
     'by repeated full-stroke cycling.'),
    ('S-5', 'Raise minimum plant load limitation if practicable',
     'Consider operational guidance to avoid sustained plant operation at load conditions '
     'where M-BFP flow remains near the minimum flow threshold for extended periods, '
     'until control logic adjustments are verified.'),
]

for code, title, desc in suggestions:
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.2)
    run_code = p.add_run(f'{code}. {title}: ')
    run_code.bold = True
    run_code.font.size = Pt(11)
    run_desc = p.add_run(desc)
    run_desc.font.size = Pt(11)

doc.add_paragraph()

# ─── 5. EXECUTION PLAN ───────────────────────────────────────────────────────
add_heading(doc, '5. EXECUTION PLAN', level=2, color=(0x1F, 0x4E, 0x79))

ep_table = doc.add_table(rows=1, cols=5)
ep_table.style = 'Table Grid'
ep_table.alignment = WD_TABLE_ALIGNMENT.CENTER
add_table_header_row(ep_table, ['No.', 'Action', 'Responsible', 'Target Date', 'Status'])

ep_data = [
    ('1', 'Review MD109 logic — check current OPEN/CLOSE setpoints for MIN FCV; '
          'calculate recommended hysteresis band with KSB/Toshiba',
     'C&I / Operations', '20 Jun 2026', 'Open'),
    ('2', 'Update DCS parameter for MIN FCV hysteresis band (MD109) after approval',
     'C&I Engineer', '30 Jun 2026', 'Open'),
    ('3', 'Apply signal damping filter on 10LAB17CF301 (MBFP Inlet Water Flow)',
     'C&I Engineer', '30 Jun 2026', 'Open'),
    ('4', 'Configure ramp rate limiter for MIN FCV (MD109) in AUTO mode',
     'C&I Engineer', '30 Jun 2026', 'Open'),
    ('5', 'Physical inspection of MIN FCV (10LAB33AA502) actuator, positioner, '
          'and valve internals during next available maintenance window',
     'Mechanical Maint.', 'Next outage', 'Planned'),
    ('6', 'Conduct function test of MIN FCV after DCS parameter changes — '
          'per VP1-C-L2-F-GEN-00029 commissioning procedure',
     'C&I / Operations', 'After action 2–4', 'Planned'),
    ('7', 'Update operating procedure for M-BFP partial load conditions; '
          'add guidance for MIN FCV manual override if oscillation observed',
     'Operations Dept.', '15 Jul 2026', 'Open'),
    ('8', 'Monitor MIN FCV position trend and MBFP flow trend for 2 weeks '
          'after DCS changes to confirm stability',
     'Operations Dept.', 'Ongoing', 'Open'),
]
for no, action, resp, date, status in ep_data:
    row = ep_table.add_row()
    row.cells[0].text = no
    row.cells[0].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    row.cells[0].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[1].text = action
    row.cells[1].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[2].text = resp
    row.cells[2].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[3].text = date
    row.cells[3].paragraphs[0].runs[0].font.size = Pt(10)
    row.cells[4].text = status
    row.cells[4].paragraphs[0].runs[0].font.size = Pt(10)
    if status == 'Open':
        set_cell_background(row.cells[4], 'FFE699')
    elif status == 'Planned':
        set_cell_background(row.cells[4], 'DDEBF7')

doc.add_paragraph()

# ─── 6. CONCLUSION ───────────────────────────────────────────────────────────
add_heading(doc, '6. CONCLUSION', level=2, color=(0x1F, 0x4E, 0x79))

p = doc.add_paragraph()
p.add_run(
    "The oscillation of the M-BFP Minimum Flow Control Valve (10LAB33AA502) was caused primarily "
    "by insufficient hysteresis in the DCS control logic (MD109) combined with flow signal "
    "instability at partial load conditions. The pump was protected from damage as the operator "
    "promptly switched the valve to manual mode.\n\n"
    "The recommended corrective actions focus on: (1) widening the control hysteresis band, "
    "(2) applying signal filtering on the flow measurement, (3) adding a valve movement ramp "
    "rate limiter, and (4) physical inspection of the valve and actuator. These measures align "
    "with requirements stated in the KSB O&M Manual (Doc: 02.05.08) and the Toshiba commissioning "
    "procedure (Doc: VP1-C-L2-F-GEN-00029).\n\n"
    "No equipment damage or loss of generation was reported. The plant continued to operate "
    "normally after the operator intervention. Follow-up monitoring and systematic DCS parameter "
    "adjustment are required to prevent recurrence."
).font.size = Pt(11)

doc.add_paragraph()

# ─── 7. ATTACHMENT ───────────────────────────────────────────────────────────
add_heading(doc, '7. ATTACHMENT', level=2, color=(0x1F, 0x4E, 0x79))

attachments = [
    ('ATT-1', 'M-BFP MIN FCV (10LAB33AA502) Position Trend — Event Period (to be attached from DCS historian)'),
    ('ATT-2', 'MBFP Inlet Water Flow (10LAB17CF301) Trend — Event Period'),
    ('ATT-3', 'MBFP Outlet Pressure Trend — Event Period'),
    ('ATT-4', 'DCS Alarm Log — MBFP MIN FCV ABN during event'),
    ('ATT-5', 'Ref: KSB O&M Manual extract — Minimum Flow Valve requirement (Doc: 02.05.08, Section 5.8.5.3 & 6.4.3)'),
    ('ATT-6', 'Ref: Logic Diagram MD109 — M-BFP MIN FCV Control Logic (Doc: VP1-C-L2-I-CB-00006, Sheet TM184/184B)'),
    ('ATT-7', 'Ref: MDBFP Data Sheet — Minimum Flow Specifications (Doc: VP1-C-L2-M-LAC-00015)'),
]
for code, desc in attachments:
    p = doc.add_paragraph(style='List Bullet')
    run_code = p.add_run(f'{code}: ')
    run_code.bold = True
    run_code.font.size = Pt(11)
    p.add_run(desc).font.size = Pt(11)

doc.add_paragraph()

# ─── SIGN-OFF ─────────────────────────────────────────────────────────────────
sign_table = doc.add_table(rows=2, cols=3)
sign_table.style = 'Table Grid'
sign_table.alignment = WD_TABLE_ALIGNMENT.CENTER
headers_sign = ['Prepared By', 'Reviewed By', 'Approved By']
for i, h in enumerate(headers_sign):
    sign_table.rows[0].cells[i].text = h
    sign_table.rows[0].cells[i].paragraphs[0].runs[0].bold = True
    sign_table.rows[0].cells[i].paragraphs[0].runs[0].font.size = Pt(10)
    sign_table.rows[0].cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_cell_background(sign_table.rows[0].cells[i], 'BDD7EE')

for i in range(3):
    cell = sign_table.rows[1].cells[i]
    p = cell.paragraphs[0]
    p.add_run('\n\n\n')
    p2 = cell.add_paragraph()
    p2.add_run('Name: _________________').font.size = Pt(9)
    p3 = cell.add_paragraph()
    p3.add_run('Date: _________________').font.size = Pt(9)

# ─── Save ─────────────────────────────────────────────────────────────────────
output_path = r"E:\book\Vanphong\tim file\7\KV-OP-26-0001_MBFP_MinimumFlowValve_Oscillation.docx"
doc.save(output_path)
print(f"Report saved to: {output_path}")
