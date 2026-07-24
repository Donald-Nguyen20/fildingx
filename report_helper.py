"""report_helper.py — Utility functions for VP1 operation reports.

Claude uses this module when generating report scripts via Bash.

Quick example:
    from report_helper import new_doc, add_title_block, add_section, add_para
    from report_helper import add_data_table, add_sign_off, add_attachment_list, save_report

    doc = new_doc()
    add_title_block(doc,
        doc_no="KV-OP-26-0002", date="13 June 2026",
        equipment="Induced Draft Fan (IDF)", unit="Unit 1",
        kks_tag="10HNA10AN001", subject="High Bearing Temperature Trip",
        prepared_by="Operations Dept.")
    add_section(doc, "1. INTRODUCTION")
    add_para(doc, "This report describes ...")
    add_section(doc, "2. ABNORMAL STATUS")
    add_data_table(doc,
        headers=["Time", "Event", "Value"],
        rows=[["T+00:00", "IDF bearing temp high", "95°C"]])
    add_sign_off(doc)
    path = save_report(doc, "KV-OP-26-0002_IDF_BearingTemp")
    print("Saved:", path)
"""
from __future__ import annotations

import datetime
import os
import sys
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

# Output directory — luôn nằm CẠNH app (cùng cấp .exe), không nằm trong
# thư mục đóng gói _internal. Khi chạy dev thì nằm cạnh source.
if getattr(sys, "frozen", False):
    _APP_BASE = Path(sys.executable).parent
else:
    _APP_BASE = Path(__file__).parent
REPORTS_DIR = _APP_BASE / "reports"
REPORTS_DIR.mkdir(exist_ok=True)


# ── Internal helpers ──────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _header_row(table, headers: list[str], bg: str = "1F4E79") -> None:
    row = table.rows[0]
    for i, h in enumerate(headers):
        cell = row.cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, bg)


# ── Public API ────────────────────────────────────────────────────────────────

def new_doc() -> Document:
    """Create Document with standard VP1 page margins."""
    doc = Document()
    for sec in doc.sections:
        sec.top_margin    = Cm(2.5)
        sec.bottom_margin = Cm(2.5)
        sec.left_margin   = Cm(3.0)
        sec.right_margin  = Cm(2.5)
    return doc


def add_title_block(
    doc: Document,
    doc_no: str,
    date: str,
    equipment: str,
    unit: str = "Unit 1",
    kks_tag: str = "",
    subject: str = "",
    discipline: str = "Operation",
    prepared_by: str = "Operations Dept.",
    rev: str = "A",
) -> None:
    """Add title header + info table (standard VP1 format)."""
    for text, size in [
        ("VAN PHONG 1 BOT THERMAL POWER PLANT", 14),
        ("OPERATION REPORT", 13),
    ]:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(size)
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)

    doc.add_paragraph()

    tbl = doc.add_table(rows=5, cols=4)
    tbl.style = "Table Grid"
    rows_data = [
        ("Document No.", doc_no,    "Date",        date),
        ("Equipment",    equipment,  "Unit",        unit),
        ("KKS Tag",      kks_tag,    "Discipline",  discipline),
        ("Subject",      subject,    "Prepared by", prepared_by),
        ("Rev.",         rev,        "Approved by", ""),
    ]
    for r_idx, (k1, v1, k2, v2) in enumerate(rows_data):
        row = tbl.rows[r_idx]
        for c_idx, val in enumerate([k1, v1, k2, v2]):
            cell = row.cells[c_idx]
            cell.text = ""
            run = cell.paragraphs[0].add_run(val)
            run.font.size = Pt(10)
            if c_idx % 2 == 0:
                run.bold = True
                _set_cell_bg(cell, "BDD7EE")

    doc.add_paragraph()


def add_section(doc: Document, title: str, level: int = 2) -> None:
    """Add a section heading (blue, VP1 style)."""
    p = doc.add_heading(title, level=level)
    for run in p.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)


def add_para(doc: Document, text: str, indent: bool = False) -> None:
    """Add a body paragraph."""
    p = doc.add_paragraph()
    if indent:
        p.paragraph_format.left_indent = Cm(0.5)
    run = p.add_run(text)
    run.font.size = Pt(11)


def add_bullet(doc: Document, text: str, bold_prefix: str = "") -> None:
    """Add a bullet point, optionally with bold prefix followed by normal text."""
    p = doc.add_paragraph(style="List Bullet")
    if bold_prefix:
        r = p.add_run(bold_prefix)
        r.bold = True
        r.font.size = Pt(11)
    p.add_run(text).font.size = Pt(11)


def add_data_table(
    doc: Document,
    headers: list[str],
    rows: list[list[str]],
    header_color: str = "1F4E79",
    alt_color: str = "",
) -> None:
    """Add a formatted data table.

    Args:
        headers:      Column headers.
        rows:         List of row value lists (strings).
        header_color: Hex background for header row.
        alt_color:    Optional alternating row background hex (e.g. 'DDEBF7').
    """
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    _header_row(tbl, headers, bg=header_color)
    for i, row_vals in enumerate(rows):
        row = tbl.add_row()
        bg = alt_color if (alt_color and i % 2 == 1) else ""
        for j, val in enumerate(row_vals):
            cell = row.cells[j]
            cell.text = ""
            run = cell.paragraphs[0].add_run(str(val))
            run.font.size = Pt(10)
            if bg:
                _set_cell_bg(cell, bg)
    doc.add_paragraph()


def add_sign_off(doc: Document) -> None:
    """Add Prepared By / Reviewed By / Approved By sign-off table."""
    tbl = doc.add_table(rows=2, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(["Prepared By", "Reviewed By", "Approved By"]):
        cell = tbl.rows[0].cells[i]
        cell.text = ""
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_cell_bg(cell, "BDD7EE")
    for i in range(3):
        cell = tbl.rows[1].cells[i]
        cell.paragraphs[0].add_run("\n\n\n")
        cell.add_paragraph().add_run("Name: _________________").font.size = Pt(9)
        cell.add_paragraph().add_run("Date: _________________").font.size = Pt(9)


def add_attachment_list(doc: Document, attachments: list[str]) -> None:
    """Add numbered attachment list.

    Args:
        attachments: List of strings like "ATT-1: DCS trend log ..."
    """
    for item in attachments:
        if ": " in item:
            code, desc = item.split(": ", 1)
            add_bullet(doc, desc, bold_prefix=f"{code}: ")
        else:
            add_bullet(doc, item)


def save_report(doc: Document, filename: str) -> str:
    """Save report to REPORTS_DIR.

    Args:
        filename: File name without extension, e.g. 'KV-OP-26-0002_IDF_BearingTemp'.
    Returns:
        Absolute path of saved file.
    """
    if not filename.endswith(".docx"):
        filename += ".docx"
    path = REPORTS_DIR / filename
    doc.save(str(path))
    return str(path)
