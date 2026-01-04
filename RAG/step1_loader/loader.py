import os
from typing import List, Iterable
from .schema import RAGDocument, LoadItemResult
from .text_cleaner import clean_vietnamese_text


class Step1Loader:
    """
    STEP 1:
    Load file -> RAGDocument(s)
    - PDF: page-level (LangChain PyPDFLoader)
    - TXT/MD/DOCX: document-level
    """

    def load_many(self, paths: Iterable[str]) -> List[LoadItemResult]:
        out = []
        for p in paths:
            out.extend(self.load_one(p))
        return out

    def load_one(self, path: str) -> List[LoadItemResult]:
        if not os.path.exists(path):
            return [LoadItemResult(False, error=f"File not found: {path}")]

        ext = os.path.splitext(path)[1].lower()

        if ext == ".pdf":
            return self._load_pdf(path)
        if ext in [".txt", ".md", ".log", ".csv"]:
            return [self._load_text(path)]
        if ext == ".docx":
            return [self._load_docx(path)]

        return [LoadItemResult(False, error=f"Unsupported file type: {ext}")]

    def _base_meta(self, path: str):
        st = os.stat(path)
        return {
            "source_path": os.path.abspath(path),
            "file_name": os.path.basename(path),
            "mtime": int(st.st_mtime),
            "size": st.st_size,
        }

    def _load_text(self, path: str) -> LoadItemResult:
        try:
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                text = clean_vietnamese_text(f.read())
            if not text:
                return LoadItemResult(False, error="Empty text")
            return LoadItemResult(True, RAGDocument(text, {**self._base_meta(path), "type": "text"}))
        except Exception as e:
            return LoadItemResult(False, error=str(e))

    def _load_docx(self, path: str) -> LoadItemResult:
        try:
            from docx import Document
            d = Document(path)
            text = clean_vietnamese_text("\n".join(p.text for p in d.paragraphs if p.text))
            return LoadItemResult(True, RAGDocument(text, {**self._base_meta(path), "type": "docx"}))
        except Exception as e:
            return LoadItemResult(False, error=str(e))

    def _load_pdf(self, path: str) -> List[LoadItemResult]:
        try:
            from langchain_community.document_loaders import PyPDFLoader
            loader = PyPDFLoader(path)
            docs = loader.load()

            results = []
            for d in docs:
                txt = clean_vietnamese_text(d.page_content)
                if not txt:
                    continue
                meta = {
                    **self._base_meta(path),
                    "type": "pdf_page",
                    "page": d.metadata.get("page", 0) + 1
                }
                results.append(LoadItemResult(True, RAGDocument(txt, meta)))

            return results or [LoadItemResult(False, error="No text in PDF")]
        except Exception as e:
            return [LoadItemResult(False, error=str(e))]
